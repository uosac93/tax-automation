"""
법인세 신고 검토 보고서 Word(docx) 생성 - 바탕체
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_NAME = "바탕"
NAVY = RGBColor(0x1a, 0x28, 0x47)
BLUE = RGBColor(0x2a, 0x5a, 0x8c)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1a, 0x1a, 0x1a)
GRAY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xCC, 0x33, 0x33)
GREEN = RGBColor(0x2a, 0x7a, 0x3a)
LIGHT_BG = "E8ECF0"
ACCENT_BG = "D6E4F0"
HDR_BG = "1A2847"


def _fmt(val):
    if val is None:
        return "-"
    if isinstance(val, (int, float)):
        v = int(val)
        if v < 0:
            return f"△{abs(v):,}"
        return f"{v:,}"
    return str(val)


def _pct(val, total):
    if not total or not val:
        return "-"
    return f"{val / total * 100:.1f}%"


def _delta(cur, prev):
    if cur is None or prev is None:
        return "-"
    d = cur - prev
    if d > 0:
        return f"+{d:,}"
    elif d < 0:
        return f"-{abs(d):,}"
    return "0"


def _delta_pct(cur, prev):
    if cur is None or prev is None or prev == 0:
        return "-"
    d = (cur - prev) / abs(prev) * 100
    return f"{d:+.1f}%"


def _set_cell_font(cell, text, size=9, bold=False, color=BLACK, align="left", font_name=FONT_NAME):
    """셀에 텍스트 + 서식 설정"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    # 셀 패딩
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 상하 패딩 줄이기
    for attr, val in [("top", "40"), ("bottom", "40"), ("left", "60"), ("right", "60")]:
        mar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:{attr} w:w="{val}" w:type="dxa"/></w:tcMar>')
        # 기존 것 제거 후 추가
    return run


def _shade_cell(cell, color_hex):
    """셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _shade_row(row, color_hex):
    """행 전체 배경색"""
    for cell in row.cells:
        _shade_cell(cell, color_hex)


def _set_table_borders(table):
    """표에 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def generate_report_docx(data, results, output_path):
    """법인세 검토 보고서 Word 생성"""
    info = data.get("회사정보", {})
    tax = data.get("세액조정", {})
    fs = data.get("재무제표", {})
    is_data = fs.get("손익계산서", {})
    is_prev = fs.get("손익계산서_전기", {})
    if not is_prev:
        is_prev = data.get("손익계산서_전기", {})

    raw_name = info.get("법인명", "")
    import re as _re
    _m = _re.match(r'^주식회사\s*(.+)$', raw_name)
    if _m:
        corp_name = f"(주){_m.group(1)}"
    else:
        _m = _re.match(r'^(.+?)\s*주식회사$', raw_name)
        corp_name = f"{_m.group(1)}(주)" if _m else raw_name
    period_start = info.get("사업연도_시작", "")
    period_end = info.get("사업연도_종료", "")
    biz_no = info.get("사업자등록번호", "")
    cur_year = period_end[:4] if period_end else "당기"
    prev_year = str(int(cur_year) - 1) if cur_year.isdigit() else "전기"
    sales_cur = is_data.get("매출액")
    sales_prev = is_prev.get("매출액")

    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # ═══════════════════════════════════════
    #  헤더
    # ═══════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{corp_name} 법인세 신고 사항")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    title.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = corp_name
    if biz_no:
        info_text += f"  |  {biz_no}"
    info_text += f"  |  {period_start} ~ {period_end}"
    run = sub.add_run(info_text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    sub.space_after = Pt(2)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_p.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    date_p.space_after = Pt(8)

    # 구분선
    doc.add_paragraph().add_run().font.size = Pt(1)

    # ═══════════════════════════════════════
    #  01. 손익계산서 주요 항목 비교
    # ═══════════════════════════════════════
    _add_section_title(doc, "01", "손익계산서 주요 항목 비교")

    # 단위 표시
    unit_p = doc.add_paragraph()
    unit_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = unit_p.add_run("(단위: 원)")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(7.5)
    run.font.color.rgb = GRAY
    unit_p.space_after = Pt(2)

    cols = 7
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # 헤더
    headers = ["계정과목", f"{cur_year}년(당기)", "비율", f"{prev_year}년(전기)", "비율", "증감액", "증감률"]
    hdr_row = table.rows[0]
    _shade_row(hdr_row, HDR_BG)
    for i, txt in enumerate(headers):
        _set_cell_font(hdr_row.cells[i], txt, 8.5, True, WHITE, "center")

    # 데이터
    items = [
        ("매출액", True), ("매출원가", False), ("매출총이익", True),
        ("판관비", False), ("영업이익", True),
        ("영업외수익", False), ("영업외비용", False),
        ("법인세차감전이익", True), ("법인세등", False), ("당기순이익", True),
    ]

    for idx, (key, is_bold) in enumerate(items):
        cur = is_data.get(key)
        prev = is_prev.get(key)
        dv = _delta(cur, prev)
        dp = _delta_pct(cur, prev)

        row = table.add_row()
        bg = ACCENT_BG if is_bold else (LIGHT_BG if idx % 2 == 0 else "FFFFFF")
        _shade_row(row, bg)

        name_color = NAVY if is_bold else BLACK
        _set_cell_font(row.cells[0], key, 9, is_bold, name_color, "left")
        _set_cell_font(row.cells[1], _fmt(cur), 9, is_bold, BLACK, "right")
        _set_cell_font(row.cells[2], _pct(cur, sales_cur), 8, False, GRAY, "right")
        _set_cell_font(row.cells[3], _fmt(prev), 9, is_bold, BLACK, "right")
        _set_cell_font(row.cells[4], _pct(prev, sales_prev), 8, False, GRAY, "right")

        delta_color = GREEN if dv.startswith("+") else (RED if dv.startswith("-") and dv != "-" else BLACK)
        _set_cell_font(row.cells[5], dv, 9, is_bold, delta_color, "right")
        dpct_color = GREEN if dp.startswith("+") else (RED if dp.startswith("-") and dp != "-" else GRAY)
        _set_cell_font(row.cells[6], dp, 8, False, dpct_color, "right")

    # 열 너비 설정
    widths = [Cm(3.2), Cm(3.0), Cm(1.5), Cm(3.0), Cm(1.5), Cm(3.0), Cm(1.8)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.add_paragraph().space_after = Pt(8)

    # ═══════════════════════════════════════
    #  02. 주요 세무조정 내역
    # ═══════════════════════════════════════
    _add_section_title(doc, "02", "주요 세무조정 내역")

    adj = data.get("소득금액조정", {})

    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table2)

    hdr2 = table2.rows[0]
    _shade_row(hdr2, HDR_BG)
    for i, txt in enumerate(["구분", "조정항목", "금액", "소득처분"]):
        _set_cell_font(hdr2.cells[i], txt, 8.5, True, WHITE, "center")

    r = 0
    for prefix, items_key in [("익금산입", "익금산입_항목"), ("손금산입", "손금산입_항목")]:
        for item in adj.get(items_key, []):
            row = table2.add_row()
            bg = LIGHT_BG if r % 2 == 0 else "FFFFFF"
            _shade_row(row, bg)
            항목 = item.get("항목명", "") or item.get("과목", "")
            _set_cell_font(row.cells[0], prefix, 9, False, BLUE, "center")
            _set_cell_font(row.cells[1], 항목, 9, False, BLACK, "left")
            _set_cell_font(row.cells[2], _fmt(item.get("금액")), 9, False, BLACK, "right")
            _set_cell_font(row.cells[3], item.get("처분", ""), 9, False, GRAY, "center")
            r += 1

    # 합계
    for lbl, key in [("익금산입 합계", "익금산입_합계"), ("손금산입 합계", "손금산입_합계")]:
        val = adj.get(key)
        if val:
            row = table2.add_row()
            _shade_row(row, ACCENT_BG)
            _set_cell_font(row.cells[0], "", 9, False, BLACK, "center")
            _set_cell_font(row.cells[1], lbl, 9, True, NAVY, "left")
            _set_cell_font(row.cells[2], _fmt(val), 9, True, NAVY, "right")
            _set_cell_font(row.cells[3], "", 9, False, BLACK, "center")

    widths2 = [Cm(3.0), Cm(6.5), Cm(4.0), Cm(3.5)]
    for row in table2.rows:
        for i, w in enumerate(widths2):
            row.cells[i].width = w

    doc.add_paragraph().space_after = Pt(8)

    # ═══════════════════════════════════════
    #  03. 법인세 산출 내역
    # ═══════════════════════════════════════
    _add_section_title(doc, "03", "법인세 산출 내역")

    납부 = tax.get("차감납부할세액")
    지방세 = int(납부 * 0.1) if 납부 else None
    합계납부 = (납부 + 지방세) if (납부 and 지방세) else None

    tax_items = [
        ("결산서상 당기순손익", tax.get("결산서상당기순손익"), False),
        ("(+) 익금산입", tax.get("익금산입"), False),
        ("(-) 손금산입", tax.get("손금산입"), False),
        ("각사업연도 소득금액", tax.get("각사업연도소득금액"), True),
        ("(-) 이월결손금 공제", tax.get("이월결손금공제"), False),
        ("과세표준", tax.get("과세표준"), True),
        ("산출세액", tax.get("산출세액"), False),
        ("(-) 공제감면세액", tax.get("최저한세적용대상_공제감면세액"), False),
        ("법인세 차감납부세액", 납부, True),
        ("법인지방소득세 (10%)", 지방세, False),
        ("합계 납부세액", 합계납부, True),
    ]

    table3 = doc.add_table(rows=1, cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table3)

    hdr3 = table3.rows[0]
    _shade_row(hdr3, HDR_BG)
    _set_cell_font(hdr3.cells[0], "구분", 8.5, True, WHITE, "center")
    _set_cell_font(hdr3.cells[1], "금액 (원)", 8.5, True, WHITE, "center")

    r = 0
    for label, val, is_bold in tax_items:
        if val is None and not is_bold:
            continue
        row = table3.add_row()
        bg = ACCENT_BG if is_bold else (LIGHT_BG if r % 2 == 0 else "FFFFFF")
        _shade_row(row, bg)
        co = NAVY if is_bold else BLACK
        _set_cell_font(row.cells[0], label, 9, is_bold, co, "left")
        _set_cell_font(row.cells[1], _fmt(val), 9, is_bold, co, "right")
        r += 1

    widths3 = [Cm(8.5), Cm(8.5)]
    for row in table3.rows:
        for i, w in enumerate(widths3):
            row.cells[i].width = w

    # ═══════════════════════════════════════
    #  푸터
    # ═══════════════════════════════════════
    doc.add_paragraph().space_after = Pt(16)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"Corp Tax_AI  |  {datetime.now().strftime('%Y.%m.%d %H:%M')}  |  박양훈 세무사")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(7.5)
    run.font.color.rgb = GRAY

    doc.save(output_path)
    return output_path


def _add_section_title(doc, num, title):
    """섹션 제목"""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)

    run_num = p.add_run(f"  {num}  ")
    run_num.font.name = FONT_NAME
    run_num._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run_num.font.size = Pt(9)
    run_num.font.bold = True
    run_num.font.color.rgb = WHITE
    # 번호 배경색
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{HDR_BG}"/>')
    run_num._element.get_or_add_rPr().append(shading)

    run_title = p.add_run(f"  {title}")
    run_title.font.name = FONT_NAME
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run_title.font.size = Pt(12)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY
